"""
This module contains functions for creating Word documents from data.
It includes functionality to create individual documents for each data row and a summary document with a table.
"""

import base64
import os
from io import BytesIO

from PIL import Image
from docx import Document
from docx.shared import Inches

from constants import COLUMNS, TITLE, IMAGE_WIDTH, IMAGE_HEIGHT, ColumnType

def create_doc(row, folder_path):
    """
    Create a Word document for a row in the data frame and save it to the specified folder.
    Args:
        row (pd.Series): The row in the data frame.
        folder_path (str): The path to the folder where the document will be saved.
    """
    doc = Document()
    doc.add_heading(TITLE, 0)

    for column in COLUMNS:
        if column.column_type == ColumnType.STRING:
            doc.add_paragraph(f"{column.display_name}: {row[column.name]}")
        elif column.column_type == ColumnType.IMAGE:
            try:
                base64_signature = row[column.name]
                base64_signature = base64_signature.replace("data:image/png;base64,", "")
                signature_bytes = base64.b64decode(base64_signature)
                img = Image.open(BytesIO(signature_bytes))

                temp_signature_path = "temp_signature.png"
                img.save(temp_signature_path)

                doc.add_picture(temp_signature_path, width=Inches(IMAGE_WIDTH), height=Inches(IMAGE_HEIGHT))

                os.remove(temp_signature_path)
            except Exception as e:
                print(e)
                doc.add_paragraph(f"{column.display_name}: {row[column.name]}")

    doc.save(f"{folder_path}/{row['num']}.docx")


def create_doc_with_table(df, folder_path):
    """
    Create a single Word document with a table containing all rows from the dataframe.
    Args:
        df (pd.DataFrame): The dataframe containing all rows.
        folder_path (str): The path to the folder where the document will be saved.
    """
    doc = Document()
    doc.add_heading(TITLE, 0)

    table = doc.add_table(rows=1, cols=len(COLUMNS))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for idx, column in enumerate(COLUMNS):
        header_cells[idx].text = column.display_name

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for idx, column in enumerate(COLUMNS):
            if column.column_type == ColumnType.STRING:
                row_cells[idx].text = str(row[column.name])
            elif column.column_type == ColumnType.IMAGE:
                try:
                    base64_signature = row[column.name]
                    base64_signature = base64_signature.replace("data:image/png;base64,", "")
                    signature_bytes = base64.b64decode(base64_signature)
                    img = Image.open(BytesIO(signature_bytes))

                    temp_signature_path = "temp_signature.png"
                    img.save(temp_signature_path)

                    paragraph = row_cells[idx].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(temp_signature_path, width=Inches(2.0), height=Inches(2.0))

                    os.remove(temp_signature_path)
                except Exception as e:
                    print(e)
                    row_cells[idx].text = str(row[column.name])

    doc.save(f"{folder_path}/summary.docx")