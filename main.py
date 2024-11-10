"""
This script reads a CSV file containing
data and creates a Word document for each row in the CSV file.
"""
import base64
import os
from collections import namedtuple
from io import BytesIO
import pandas as pd
from tqdm import tqdm
from docx import Document
from docx.shared import Inches
from PIL import Image

# Define a class to represent a column in the data

Column = namedtuple("Column", ["name", "display_name", "column_type"])

class ColumnType:
    """
    An enumeration of column types.
    """
    STRING = "string"
    IMAGE = "image"

# Define the column that contains the unique identifier for each document
ID_COL = "num"

# Define the columns in the data
COLUMNS = [
    Column("block/payment_phone.text1", "שם", "string"),
    # Column("תעודת זהות", "תעודת זהות", "string"),
    # Column("סכום", "סכום", "string"),
    # Column("responseDate", "תאריך", "string"),
    Column("sign", "חתימה", "image")
]

TITLE = "חתימת נבדק" # Title of the document

# Define the size of the image to be inserted into the document
IMAGE_WIDTH = 4.0 # Inches
IMAGE_HEIGHT = 4.0 # Inches

# Define the folder where the results will be saved
RESULTS_FOLDER = "results"

# Create the results folder if it does not exist
if not os.path.exists(RESULTS_FOLDER):
    os.makedirs(RESULTS_FOLDER)

# Function to create a document for a row in the data frame and save it 
def create_doc(row, folder_path):
    """
    Create a Word document for a row in the data frame and save it to the specified folder.
    Args:
        row (pd.Series): The row in the data frame.
        folder_path (str): The path to the folder where the document will be saved.
    """
    # Create a new document
    doc = Document()

    # Add the title to the document
    doc.add_heading(TITLE, 0)


    for column in COLUMNS:
        if column.column_type == ColumnType.STRING:
            doc.add_paragraph(f"{column.display_name}: {row[column.name]}")
        elif column.column_type == ColumnType.IMAGE:
            try:
                base64_signature = row[column.name]
                base64_signature = base64_signature.replace("data:image/png;base64,", "")
                signature_bytes = base64.b64decode(base64_signature)
                img = Image.open(BytesIO(signature_bytes))

                temp_signature_path = "temp_signature.png"
                img.save(temp_signature_path)

                # Add the image to the document
                doc.add_picture(temp_signature_path, width=Inches(IMAGE_WIDTH), height=Inches(IMAGE_HEIGHT))

                # Clean up: remove the temporary signature image file
                os.remove(temp_signature_path)
            except Exception as e:
                print(e)
                doc.add_paragraph(f"{column.display_name}: {row[column.name]}")

    # Save the document to a file with the ID of the row as the filename in the results folder 
    doc.save(f"{folder_path}/{row[ID_COL]}.docx")

if __name__ == "__main__":
    FILE_NAME = './payment.csv'  # The name of the CSV file

    # Load the data from the CSV file into a pandas DataFrame
    df = pd.read_csv(FILE_NAME)  # Load the data

    # Iterate over the rows in the DataFrame and create a document for each row
    for index, r in tqdm(df.iterrows(), total=df.shape[0]):
        create_doc(r, RESULTS_FOLDER)

    # Open the results folder in the file explorer
    os.startfile(RESULTS_FOLDER)
